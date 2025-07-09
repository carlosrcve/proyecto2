from django.shortcuts import render, redirect
from django.http import HttpResponse,JsonResponse
from django.http.response import HttpResponse
from django.contrib import messages
from django.core.files import File
from .models import *
import datetime
import hashlib
#from . models import Post
from openpyxl import Workbook  # Para generar el informe en excel
from django.core.exceptions import ObjectDoesNotExist
import mimetypes
import openpyxl
from openpyxl.styles import Font, PatternFill,NamedStyle,Alignment
from openpyxl.styles.borders import Border, Side
import openpyxl.styles as os
import os
from django.http import FileResponse
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO

# Create your views here.



def homePage(request):
    return render(request,"App/index.html")

# Ruta
def descargar_excel_ruta(request):
    # Reemplaza con tu URL de GitHub
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
    filename = 'data_center.xls'
    filepath = BASE_DIR + '/App/archivos_excel/' + filename 
    path = open(filepath, 'r') 
    mime_type, _ = mimetypes.guess_type(filepath)


    # Abre el archivo usando la URL
    response = HttpResponse(File(open(filepath, 'rb')), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="data_center.xls"'
    return response


def descargar_pdf_ruta(request):
    # Reemplaza con tu URL de GitHub
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
    filename = 'Privado.pdf'
    filepath = BASE_DIR + '/App/archivos_pdf/' + filename 
    path = open(filepath, 'r') 
    mime_type, _ = mimetypes.guess_type(filepath)


    # Abre el archivo usando la URL
    response = HttpResponse(File(open(filepath, 'rb')), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Privado.pdf"'
    return response

def descargar_word_ruta(request):
    # Reemplaza con tu URL de GitHub
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
    filename = 'documento_word.docx'
    filepath = BASE_DIR + '/App/archivos_word/' + filename 
    path = open(filepath, 'r') 
    mime_type, _ = mimetypes.guess_type(filepath)


    # Abre el archivo usando la URL
    response = HttpResponse(File(open(filepath, 'rb')), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="documento_word.docx"'
    return response

def informe_empleado(request):
    try:
        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="indices_precios_al_consumidor.xlsx"'

        # Consulta la base de datos para obtener los datos que deseas exportar
        datos = Post.objects.all()

        # Crea un nuevo libro de Excel y una hoja de trabajo
        workbook = Workbook()
        worksheet = workbook.active

        pink = "00FF00FF"
        green = "00008000"
        black = '000000'
        thin = Side(border_style="thin", color=pink)
        double = Side(border_style="double", color=black)

        # Encabezado 1
        worksheet['A1'] = "ÍNDICE NACIONAL DE PRECIOS AL CONSUMIDOR"
        worksheet['A1'].font = Font(bold=True)
        #worksheet["A1"].border = Border(top=double, left=thin, right=thin, bottom=double)
        worksheet["A1"].border = Border(bottom=double)
        worksheet["B1"].border = Border(bottom=double)
        
        # Encabezado 2
        worksheet['A2'] = "Serie desde Diciembre  2007"
        worksheet['A2'].font = Font(bold=True)
        worksheet["A2"].border = Border(bottom=double)
        worksheet["B2"].border = Border(bottom=double)

        # Encabezado 3
        worksheet['A3'] = "( BASE Diciembre 2007 = 100 )"
        worksheet['A3'].font = Font(bold=True)
        worksheet["A3"].border = Border(bottom=double)
        worksheet["B3"].border = Border(bottom=double)

        # Definiendo el ancho de las Columnas
        worksheet.column_dimensions["A"].width = 13

        worksheet.column_dimensions["B"].width = 56

        worksheet.column_dimensions["C"].width = 8

        '''
        # Central las el contenido de als columnas
        for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row, max_col=1):  # Ajusta min_row y max_row según tu rango
            for cell in row:
                cell.alignment = openpyxl.styles.Alignment(horizontal='center')
        '''
        # Quitar la cuadricula de las rayas grises
        worksheet.sheet_view.showGridLines = False

        # Agrega encabezados
        worksheet.append(
            ['mes', 'indice', 'var'])


        # Agrega los datos a la hoja de trabajo
        for dato in datos:
            worksheet.append([dato.mes,dato.indice,dato.variacion])

        # Guarda el libro de Excel en la respuesta HTTP
        workbook.save(response)

        return response
    except ObjectDoesNotExist:
        error_message = "El Empleado con id: {id} no existe."
        return render(request, "index.html", {"error_message": error_message})


def generate_pdf(request):
    response = FileResponse(generate_pdf_file(),as_attachment=True, filename='book_catalog44.pdf')
    return response

def home(request):
    if request.method == 'POST':
        form = PostForm(request.POST)
        if form.is_valid():
            form.save()
    # Redirect to PDF generation after adding a book
            return redirect('home')  
    else:
        form = PostForm()
    return render(request, 'App/create_user_profile.html', {'form': form})

def generate_pdf_file():
    from io import BytesIO

    buffer = BytesIO()
    p = canvas.Canvas(buffer)

    # Create a PDF document
    books = Post.objects.all()
    p.drawString(100, 750, "Book Catalog")

    y = 700
    for book in books:
        p.drawString(100, y, f"Title: {book.mes}")
        p.drawString(250, y - 0, f"Author: {book.indice}")
        p.drawString(400, y - 0, f"Year: {book.variacion}")
        y -= 60

    p.showPage()
    p.save()

    buffer.seek(0)
    return buffer


def generar_informe_word(request):
    # 1. Obtener datos de la base de datos o de donde sea
    datos = Post.objects.all()  # Ejemplo con un modelo

    # 2. Crear un nuevo documento Word
    document = Document()
    document.add_heading('Informe de Datos', 0)

    # 3. Agregar contenido al documento
    for dato in datos:
        document.add_paragraph(f"Dato: {dato.mes}, Valor: {dato.indice}, Valor: {dato.variacion}")

    # 4. Guardar el documento en memoria
    document_io = BytesIO()
    document.save(document_io)
    document_io.seek(0)  # Volver al inicio del stream

    # 5. Configurar la respuesta HTTP
    response = HttpResponse(
        document_io.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=informe.docx'

    return response


def descargar_archivo_word(request):
    # Reemplaza con tu URL de GitHub
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
    filename = 'documento_word.docx'
    filepath = BASE_DIR + '/App/archivos_word/' + filename 
    path = open(filepath, 'r') 
    mime_type, _ = mimetypes.guess_type(filepath)


    # Abre el archivo usando la URL
    response = HttpResponse(File(open(filepath, 'rb')), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="documento_word.docx"'
    return response

def descargar_archivo_word1(request):
       document = Document()
       document.add_paragraph('Este es un ejemplo de texto en el documento Word.')
       # Aquí puedes agregar más contenido al documento
       
       response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
       response['Content-Disposition'] = 'attachment; filename=documento_word.docx'
       document.save(response)
       return response


def descargar_archivo_word2(request):
    # Reemplaza con tu URL de GitHub
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
    filename = 'documento_word1.docx'
    filepath = BASE_DIR + '/App/archivos_word/' + filename 
    path = open(filepath, 'r') 
    mime_type, _ = mimetypes.guess_type(filepath)


    # Abre el archivo usando la URL
    response = HttpResponse(File(open(filepath, 'rb')), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="documento_word.docx"'
    return response

def descargar_archivo_txt(request): 

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 

    filename = 'mi_archivo.txt'

    filepath = BASE_DIR + '/App/archivos/' + filename 

    path = open(filepath, 'r') 

    mime_type, _ = mimetypes.guess_type(filepath)
    
    response = HttpResponse(path, content_type = mime_type)

    response['Content-Disposition'] = f"attachment; filename={filename}"

    return response 


def prueba(request):
    return render(request,"App/carlos.html")


def bcv1(request):
    return render(request,"App/bcv.html")


def descargar_informes_financieros(request):
    return render(request,"App/estadistica.html")


